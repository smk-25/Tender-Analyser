# TenderAnalyzer_Consolidated.py
"""
Streamlit Tender Analyzer (semantic chunking + map/reduce)
- Supports multiple PDF uploads
- Includes timing: extraction, chunking, map time, reduce time (secs)
- Uses google-genai (Gemini) for map & reduce (Parallel Map Phase)
pip install streamlit pymupdf nltk google-genai
"""

import os
import re
import json
import time
import tempfile
import asyncio
import collections
import io
from typing import List, Dict, Any, Optional, Tuple

import pandas as pd
import streamlit as st
import fitz  # PyMuPDF
import nltk
from google import genai
from docx import Document
import pdfplumber

# Ensure nltk resources (quiet, cached)
@st.cache_resource
def setup_nltk():
    nltk.download('punkt', quiet=True)
    nltk.download('punkt_tab', quiet=True)

setup_nltk()

st.set_page_config(page_title="Tender Analyzer (Semantic + MapReduce)", layout="wide")

# --------------------------
# UI: sidebar config
# --------------------------
st.title("Tender Analyzer — Semantic Map/Reduce")
st.markdown("Upload one or more tender PDFs, type queries (one per line), and get structured summaries. Uses Gemini (google-genai).")

with st.sidebar:
    # API Status Indicator
    if "api_backoff_until" in st.session_state:
        backoff_time = st.session_state["api_backoff_until"]
        if time.time() < backoff_time:
            st.warning(f"⏳ API Rate Limit: Backing off for {int(backoff_time - time.time())}s...")
        else:
            st.success("✅ API Status: Ready")
            
    st.header("Settings")
    # Using unique keys to prevent DuplicateElementId if the script ever re-runs unexpectedly
    model = st.text_input("Model (google-genai)", value="gemini-2.5-flash", key="model_input")
    target_rpm = st.number_input("Target RPM (requests/min)", value=4, min_value=1, key="rpm_input")
    batch_size = st.number_input("Batch size", value=20, min_value=1, key="batch_input")
    max_tokens_per_chunk = st.number_input("Max tokens per chunk", value=8000, min_value=1000, key="max_tokens_input")
    overlap_sentences = st.number_input("Overlap sentences", value=5, min_value=0, key="overlap_input")
    
    st.markdown("---")
    analysis_mode = st.radio(
        "Select objective",
        ["General Summary", "Compliance Matrix", "Risk Assessment", "Entity Dashboard", "Ambiguity Scrutiny"],
        key="mode_radio"
    )
    
    enable_vision = st.checkbox("Enable Vision (Scrutinize charts/images)", value=False, help="Uses Multimodal Vision for image-heavy pages.")
    
    st.markdown("**API Key**: set `GOOGLE_API_KEY` env var or paste below.")
    api_key_input = st.text_input("Google API Key (optional)", type="password", key="api_key_input")
    
    st.markdown("---")
    prompt_instructions = st.text_area(
        "Optional extra prompt instructions (prepended to LLM prompt)",
        value="",
        help="Example: 'Return answer only — do not include sources.' or 'Focus on technical requirements.'",
        key="prompt_instr_input"
    )

    st.markdown("---")
    st.header("📊 Benchmarking")
    benchmark_csv = st.file_uploader("Upload Evaluation CSV (ID, queries, answer)", type=["csv"], key="benchmark_csv")
    if benchmark_csv:
        t_load_start = time.time()
        df_bench = pd.read_csv(benchmark_csv)
        t_load_end = time.time()
        st.success(f"CSV Loaded in {t_load_end - t_load_start:.4f}s")
        if all(col in df_bench.columns for col in ["ID", "queries", "answer"]):
            st.session_state["benchmark_df"] = df_bench
        else:
            st.error("CSV must contain ID, queries, and answer columns.")

# --------------------------
# Helper: configure API key
# --------------------------
API_KEY = api_key_input.strip() or os.environ.get("GOOGLE_API_KEY")
if not API_KEY:
    st.warning("No API key provided yet. Enter GOOGLE_API_KEY in sidebar or set environment variable.")
    st.stop()

# Initialize client
client = genai.Client(api_key=API_KEY)

# Create debug dir
DEBUG_DIR = os.path.join(tempfile.gettempdir(), "tender_debug_streamlit_optimized")
os.makedirs(DEBUG_DIR, exist_ok=True)

class ConcurrencyManager:
    """Manages a global semaphore that is safe across Streamlit event loop re-runs."""
    def __init__(self, value: int):
        self.value = value
        self._sem = None
        self._loop = None
    
    def get(self):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            return None # Should not happen in Streamlit async context
            
        if self._sem is None or self._loop != loop:
            self._sem = asyncio.Semaphore(self.value)
            self._loop = loop
        return self._sem

GLOBAL_CONCURRENCY = ConcurrencyManager(2) # Default to 2 for better stability

# --------------------------
# Rate control wrapper (Async)
# --------------------------
class TokenBucket:
    def __init__(self, rpm: int):
        self.spacing = 60.0 / max(1, rpm)
        self.last_call = 0.0
        # Don't create the lock in __init__ if it might be used across event loops
        self._lock = None

    @property
    def lock(self):
        # Tie the lock to the current event loop to avoid RuntimeError across Streamlit re-runs
        loop = asyncio.get_event_loop()
        if self._lock is None or self._lock._loop != loop:
            self._lock = asyncio.Lock()
        return self._lock

    async def wait(self):
        async with self.lock:
            # Use a robust clock
            now = asyncio.get_event_loop().time()
            
            # If last_call is in the past, reset it to now
            if self.last_call < now:
                self.last_call = now

            wait_time = (self.last_call + self.spacing) - now
            if wait_time > 0:
                await asyncio.sleep(min(wait_time, 60.0))
                self.last_call += self.spacing
            else:
                # Add a tiny random jitter (10-50ms) to prevent perfectly synchronized bursts
                import random
                self.last_call = now + self.spacing + (random.uniform(0.01, 0.05))

    async def report_429(self, wait_sec: float):
        """Global push-back if anyone hits a 429."""
        async with self.lock:
            now = asyncio.get_event_loop().time()
            # Push the next scheduled call to at least now + wait_sec
            self.last_call = max(self.last_call, now + wait_sec)
            # Store in session state for UI feedback (if available)
            try:
                st.session_state["api_backoff_until"] = time.time() + wait_sec
            except: pass

_rate_limiters = {}

def get_limiter(rpm: int) -> TokenBucket:
    if rpm not in _rate_limiters:
        _rate_limiters[rpm] = TokenBucket(rpm)
    return _rate_limiters[rpm]

async def call_gemini_json_async(system_msg: str, user_msg: str, model_name: str, rpm: int, max_retries=10):
    limiter = get_limiter(rpm)
    backoff = 5.0
    import random
    
    for attempt in range(max_retries):
        try:
            await limiter.wait()
            sem = GLOBAL_CONCURRENCY.get()
            async with (sem if sem else asyncio.sleep(0)): # Fallback
                # Proper system instruction handling for google-genai
                from google.genai import types
                resp = await client.aio.models.generate_content(
                    model=model_name,
                    contents=user_msg,
                    config=types.GenerateContentConfig(
                        system_instruction=system_msg,
                        response_mime_type="application/json"
                    )
                )
            text = resp.text or "{}"
            
            # JSON Cleaning
            clean_json_text = text.strip()
            if clean_json_text.startswith("```json"):
                clean_json_text = clean_json_text.replace("```json", "", 1).rsplit("```", 1)[0].strip()
            elif clean_json_text.startswith("```"):
                 clean_json_text = clean_json_text.replace("```", "", 1).rsplit("```", 1)[0].strip()
            
            try:
                return json.loads(clean_json_text)
            except Exception:
                # RegEx fallback
                m = re.search(r"\{.*\}|\[.*\]", clean_json_text, re.S)
                if m: return json.loads(m.group(0))
                return {"_raw": text}

        except Exception as e:
            msg = str(e).lower()
            # Log to debug
            with open(os.path.join(DEBUG_DIR, f"api_error_{int(time.time())}_{attempt}.txt"), "w") as f:
                f.write(msg)
            
            if "429" in msg or "resource_exhausted" in msg:
                # Extract retry delay if possible
                wait_sec = backoff + (random.random() * 2)
                if "retrydelay" in msg:
                    try:
                        # Crude parsing for 'retryDelay': '32s'
                        match = re.search(r"retrydelay':\s*'(\d+)s", msg)
                        if match: wait_sec = float(match.group(1)) + 1
                    except: pass
                
                # Signal the bucket to stop all threads
                await limiter.report_429(wait_sec)
                await asyncio.sleep(wait_sec)
                backoff = min(backoff * 2.0, 120.0) # More aggressive backoff
                continue
            
            if "400" in msg or "invalid_argument" in msg:
                return {"error": f"Invalid request: {msg}"}

            if attempt == max_retries - 1: return {"error": msg}
            await asyncio.sleep(backoff + random.random())
            backoff *= 1.5
            
    return {"error": "max_retries_exceeded"}

# --------------------------
# Extraction & chunking (semantic)
# --------------------------
def clean_text(t: str) -> str:
    t = t.replace("\x00", " ").replace("\r\n", "\n")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

def extract_pages(pdf_path: str) -> Tuple[List[str], List[bool]]:
    doc = fitz.open(pdf_path)
    pages = []
    image_flags = []
    
    # Use pdfplumber for advanced table detection
    with pdfplumber.open(pdf_path) as pl_pdf:
        for i in range(len(doc)):
            page = doc[i]
            pl_page = pl_pdf.pages[i]
            
            try:
                blocks = page.get_text("blocks")
                blocks = sorted(blocks, key=lambda b: (round(b[1],1), round(b[0],1)))
                text = "\n".join(b[4] for b in blocks if isinstance(b, (list, tuple)) and len(b) >= 5)
            except Exception:
                text = page.get_text()
            
            # Extract tables and format as Markdown
            tables = pl_page.extract_tables()
            table_text = ""
            for table in tables:
                if table:
                    # Clean the table data
                    df_t = pd.DataFrame(table).dropna(how='all').dropna(axis=1, how='all')
                    if not df_t.empty:
                        table_text += f"\n\n[TABLE DATA - PAGE {i+1}]:\n" + df_t.to_markdown(index=False) + "\n\n"
            
            if table_text:
                text += table_text
                
            has_images = len(page.get_images()) > 0
            pages.append(clean_text(text))
            image_flags.append(has_images)
            
    doc.close()
    return pages, image_flags

def looks_like_table_block(text: str) -> bool:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return False
    numeric_lines = sum(1 for ln in lines if re.search(r'\d', ln))
    numeric_ratio = numeric_lines / max(1, len(lines))
    has_delims = any('|' in ln or '\t' in ln or re.search(r'\s{2,}', ln) for ln in lines)
    return numeric_ratio > 0.25 and has_delims

def semantic_chunk_pages(pages: List[str], image_flags: List[bool], model_name: str, max_tokens=8000, overlap_sentences=5):
    # Extract blocks
    flat = []
    for idx, (content, has_img) in enumerate(zip(pages, image_flags), start=1):
        content = content.strip()
        if not content: continue
        paragraphs = re.split(r'\n\s*\n', content)
        for par in paragraphs:
            par = par.strip()
            if looks_like_table_block(par):
                rows = [f"[Page {idx}] {r.strip()}" for r in par.splitlines() if r.strip()]
                for r in rows: flat.append((idx, r, has_img))
            else:
                try: sents = nltk.tokenize.sent_tokenize(par)
                except: sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', par) if s.strip()]
                if sents: sents[0] = f"[Page {idx}] " + sents[0]
                for s in sents: flat.append((idx, s, has_img))

    chunks = []
    buf_sents = []
    buf_pages = set()
    has_visual = False
    cid = 1
    
    def get_token_count(text):
        return len(text) // 4 

    for (pg, sent, img_flag) in flat:
        curr_text = " ".join(buf_sents + [sent])
        if get_token_count(curr_text) > max_tokens and buf_sents:
            text = " ".join(buf_sents).strip()
            chunks.append({
                "id": f"chunk_{cid}", 
                "text": text, 
                "start_page": min(buf_pages), 
                "end_page": max(buf_pages),
                "has_visual": has_visual
            })
            cid += 1
            has_visual = False # Reset for next chunk
            buf_sents = buf_sents[-overlap_sentences:] if overlap_sentences > 0 else []
            buf_pages = {pg}
        
        buf_sents.append(sent)
        buf_pages.add(pg)
        if img_flag: has_visual = True

    if buf_sents:
        chunks.append({
            "id": f"chunk_{cid}", 
            "text": " ".join(buf_sents).strip(), 
            "start_page": min(buf_pages), 
            "end_page": max(buf_pages),
            "has_visual": has_visual
        })
    
    return [c for c in chunks if c["text"].strip()]

# --------------------------
# Evaluation Metrics
# --------------------------
def normalize_text(s: str) -> str:
    if s is None: return ""
    s = s.lower()
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def tokenize_simple(s: str) -> List[str]:
    return re.findall(r"\w+", (s or "").lower())

def token_precision(pred: str, gold: str) -> float:
    p_tokens = tokenize_simple(pred)
    g_tokens = tokenize_simple(gold)
    if not p_tokens: return 0.0
    common = collections.Counter(p_tokens) & collections.Counter(g_tokens)
    num_same = sum(common.values())
    return num_same / len(p_tokens)

def token_recall(pred: str, gold: str) -> float:
    p_tokens = tokenize_simple(pred)
    g_tokens = tokenize_simple(gold)
    if not g_tokens: return 1.0
    common = collections.Counter(p_tokens) & collections.Counter(g_tokens)
    num_same = sum(common.values())
    return num_same / len(g_tokens)

def rouge_l(pred: str, gold: str) -> float:
    p, g = tokenize_simple(pred), tokenize_simple(gold)
    if not p or not g: return 0.0
    n, m = len(p), len(g)
    prev = [0] * (m + 1)
    for i in range(1, n + 1):
        curr = [0] * (m + 1)
        for j in range(1, m + 1):
            if p[i-1] == g[j-1]: curr[j] = prev[j-1] + 1
            else: curr[j] = max(prev[j], curr[j-1])
        prev = curr
    lcs = prev[m]
    prec, rec = lcs/n, lcs/m
    return 2*prec*rec/(prec+rec) if (prec+rec) > 0 else 0.0

# --------------------------
# Prompts & map/reduce (UI-friendly)
# --------------------------
# --------------------------
# Prompts & Map/Reduce (Async)
# --------------------------
MAP_SYS = "You are an experienced tenders analyst. Extract only content relevant to the query and cite exact PDF page numbers."
MAP_INS = (
    "You will receive multiple PDF chunks. For EACH chunk, return one JSON object:\n"
    "{\n  \"id\": \"<chunk_id>\",\n  \"bullets\": [\"short points relevant to the query\"],\n  \"citations\": [page_numbers]\n}\n"
    "Return a JSON array of these objects. Only JSON."
)
REDUCE_SYS = "You are compiling a concise answer from partial notes. Merge, deduplicate, and keep page citations."
REDUCE_INS = (
    "Given an array of items with fields id, bullets, citations, produce a single JSON object:\n"
    "{\n"
    "  \"summary\": \"2-4 short paragraphs\",\n"
    "  \"bullets\": [\"final deduped bullets\"],\n"
    "  \"citations\": [page_numbers]\n"
    "}\n"
    "Only JSON."
)
def get_prompts(mode: str):
    if mode == "Compliance Matrix":
        return ("Extract mandatory tender requirements (Technical, Financial, Legal).",
                "Return JSON array: [{\"item\": \"short name\", \"detail\": \"requirement description\", \"evidence\": \"quote from text\", \"page\": 1}]",
                "Merge and deduplicate these compliance requirements into a single master matrix.",
                "Return single JSON: {\"matrix\": [{\"item\": \"...\", \"detail\": \"...\", \"evidence\": \"...\", \"citations\": [pages]}]}")
    elif mode == "Risk Assessment":
        return ("Flag high-risk tender clauses (Liabilities, Penalties, Termination).",
                "Return JSON array: [{\"clause\": \"topic\", \"risk_level\": \"High/Medium\", \"reason\": \"why it is risky\", \"evidence\": \"quote from text\", \"page\": 1}]",
                "Merge and deduplicate these risk findings into a single consolidated risk report.",
                "Return single JSON: {\"risks\": [{\"clause\": \"...\", \"risk_level\": \"...\", \"reason\": \"...\", \"mitigation\": \"...\", \"evidence\": \"...\", \"citations\": [pages]}]}")
    elif mode == "Entity Dashboard":
         return ("Extract critical tender metadata: Deadlines, Financials (EMD/PG), Locations, and Key Contacts.",
                "Return JSON array: [{\"category\": \"Deadline/Financial/Location/Contact\", \"entity\": \"...\", \"value\": \"...\", \"page\": 1}]",
                "Compile an Executive Dashboard of key tender entities. Group by category.",
                "Return single JSON: {\"dashboard\": {\"Deadlines\": [], \"Financials\": [], \"Locations\": [], \"Contacts\": []}}")
    elif mode == "Ambiguity Scrutiny":
         return ("IDentify ambiguous, conflicting, or vague clauses that need clarification from the authority.",
                "Return JSON array: [{\"clause\": \"...\", \"issue\": \"vague/conflicting/...\", \"suggested_query\": \"formal question for bidder meet\", \"page\": 1}]",
                "Consolidate a Clarification Request document (Pre-bid queries).",
                "Return single JSON: {\"queries\": [{\"item\": \"...\", \"conflict\": \"...\", \"query\": \"...\", \"citations\": [pages]}]}")
    else:
        # Standardized General Summary prompts - More inclusive language to avoid '0 snippets'
        return ("Extract key details, scope, deliverables, or any informative sections relevant to the query. Be descriptive.",
                "Return a JSON array of findings: [{\"finding\": \"topic title\", \"snippet\": \"important details or exact quote\", \"page\": 1}]. If a chunk is purely introductory, capture its main topic anyway.",
                "Synthesize a final executive summary from these findings.",
                "Return single JSON: {\"summary\": \"...\", \"bullets\": [\"key points\"], \"citations\": [pages]}")

def generate_clarification_letter(results: List[Dict]):
    """Generate a formal Word doc for pre-bid queries if Ambiguity Scrutiny was used."""
    doc = Document()
    doc.add_heading('Pre-bid Clarification Request', 0)
    
    doc.add_paragraph('To the Tendering Authority,')
    doc.add_paragraph('\nWe are seeking clarification on the following points identified during our technical review of the tender documents:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reference/Clause'
    hdr_cells[1].text = 'Issue/Ambiguity'
    hdr_cells[2].text = 'Proposed Query'

    count = 0
    for r in results:
        # Check both manual results and benchmark results format
        res_data = r.get("result", {})
        if r.get('mode') == "Ambiguity Scrutiny":
            items = res_data.get('queries', [])
            for q_item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(q_item.get('item', 'N/A'))
                row_cells[1].text = str(q_item.get('conflict', 'Ambiguity detected'))
                row_cells[2].text = str(q_item.get('query', 'N/A'))
                count += 1
    
    if count == 0:
        return None
        
    doc.add_paragraph('\nWe look forward to your response.')
    doc.add_paragraph('\nSincerely,')
    doc.add_paragraph('Interested Bidder')
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_excel_report(results: List[Dict]):
    """Generate a multi-sheet Excel file for different analysis modes."""
    bio = io.BytesIO()
    try:
        with pd.ExcelWriter(bio, engine='openpyxl') as writer:
            # 1. Summary Sheet
            summary_rows = []
            for r in results:
                res = r.get("result", {})
                summary_rows.append({
                    "Query": r.get("query"),
                    "Mode": r.get("mode"),
                    "Summary": res.get("summary", "N/A"),
                    "Key_Points": ", ".join(res.get("bullets", []))
                })
            if summary_rows:
                pd.DataFrame(summary_rows).to_excel(writer, sheet_name="General_Summary", index=False)
            
            # 2. Compliance Sheet
            compliance_rows = []
            for r in results:
                if r.get("mode") == "Compliance Matrix":
                    for item in r.get("result", {}).get("matrix", []):
                        compliance_rows.append(item)
            if compliance_rows:
                pd.DataFrame(compliance_rows).to_excel(writer, sheet_name="Compliance_Matrix", index=False)
            
            # 3. Risks Sheet
            risk_rows = []
            for r in results:
                if r.get("mode") == "Risk Assessment":
                    for risk in r.get("result", {}).get("risks", []):
                        risk_rows.append(risk)
            if risk_rows:
                pd.DataFrame(risk_rows).to_excel(writer, sheet_name="Risk_Assessment", index=False)

            # 4. Entities Sheet
            entity_rows = []
            for r in results:
                if r.get("mode") == "Entity Dashboard":
                    dash = r.get("result", {}).get("dashboard", {})
                    for cat, items in dash.items():
                        for it in items:
                            entity_rows.append({"Category": cat, **it})
            if entity_rows:
                pd.DataFrame(entity_rows).to_excel(writer, sheet_name="Key_Entities", index=False)
    except Exception as e:
        st.error(f"Error generating Excel: {e}")
        return None

    bio.seek(0)
    return bio

# --------------------------
# PDF Utilities for Vision & Citations
# --------------------------

def get_page_image(doc: fitz.Document, page_num: int, highlight_text: str = None):
    """Render a PDF page to a bytes image with optional text highlighting."""
    try:
        page = doc.load_page(page_num - 1)
        if highlight_text:
            # Clean text for search (PyMuPDF search is sensitive)
            search_text = highlight_text.strip()
            # If text is too long, take the first 60 chars to ensure a match
            if len(search_text) > 60:
                search_text = search_text[:60]
            
            quads = page.search_for(search_text)
            for quad in quads:
                # Add a semi-transparent highlight
                annot = page.add_highlight_annot(quad)
                annot.set_colors(stroke=(1, 1, 0)) # Yellow
                annot.update(opacity=0.35)
                
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5)) # 1.5x zoom
        return pix.tobytes("png")
    except Exception as e:
        print(f"Highlight error: {e}")
        return None

def render_citation_preview(doc, citations):
    """
    Render page previews with highlights.
    citations: List[Dict] with {"page": int, "text": str}
    """
    if not citations: return
    
    st.markdown("---")
    st.markdown("### 📄 Context Preview")
    
    # Group by page to avoid redundant tabs
    from collections import defaultdict
    grouped = defaultdict(list)
    for c in citations:
        if isinstance(c, dict) and "page" in c:
            grouped[c["page"]].append(c.get("text", ""))
        elif isinstance(c, int): # Fallback
            grouped[c].append("")
    
    sorted_pages = sorted(grouped.keys())
    # Limit number of tabs to prevent UI lag
    display_pages = sorted_pages[:8]
    
    tabs = st.tabs([f"Page {p}" for p in display_pages])
    for i, p in enumerate(display_pages):
        with tabs[i]:
            # Use the first snack available for that page for highlighting
            snippets = [s for s in grouped[p] if s]
            text_to_highlight = snippets[0] if snippets else None
            
            img = get_page_image(doc, p, highlight_text=text_to_highlight)
            if img: st.image(img, caption=f"Original Document - Page {p}")
            else: st.error(f"Could not render page {p}")

async def call_gemini_vision_async(system_msg: str, user_msg: str, image_bytes: bytes, model_name: str, rpm: int, max_retries=5):
    """Handle multimodal inputs for pages with diagrams/charts with retries."""
    limiter = get_limiter(rpm)
    backoff = 5.0
    import random
    
    for attempt in range(max_retries):
        try:
            await limiter.wait()
            sem = GLOBAL_CONCURRENCY.get()
            async with (sem if sem else asyncio.sleep(0)):
                from google.genai import types
                img_part = types.Part.from_bytes(data=image_bytes, mime_type="image/png")
                
                resp = await client.aio.models.generate_content(
                    model=model_name,
                    contents=[user_msg, img_part],
                    config=types.GenerateContentConfig(
                        system_instruction=system_msg
                    )
                )
            return resp.text or ""
        except Exception as e:
            msg = str(e).lower()
            if "429" in msg or "resource_exhausted" in msg:
                wait_sec = backoff + (random.random() * 2)
                if "retrydelay" in msg:
                    try:
                        match = re.search(r"retrydelay':\s*'(\d+)s", msg)
                        if match: wait_sec = float(match.group(1)) + 1
                    except: pass
                
                await limiter.report_429(wait_sec)
                await asyncio.sleep(wait_sec)
                backoff = min(backoff * 2.0, 120.0)
                continue
            
            if attempt == max_retries - 1:
                return f"[Vision Error: {e}]"
            await asyncio.sleep(backoff + random.random())
            backoff *= 1.5

    return "[Vision Error: Max retries exceeded]"

async def map_phase_async(chunks: List[Dict], query: str, model: str, rpm: int, batch_size: int, mode: str, doc: fitz.Document = None, enable_vision: bool = False):
    m_sys, m_ins, _, _ = get_prompts(mode)
    batches = [chunks[i:i+batch_size] for i in range(0, len(chunks), batch_size)]
    
    async def process_batch(batch):
        # Handle Vision integration
        vision_info = ""
        if enable_vision and doc:
            for c in batch:
                if c.get("has_visual"):
                    img_bytes = get_page_image(doc, c["start_page"])
                    if img_bytes:
                        desc = await call_gemini_vision_async(
                            "Describe any charts, tables, or diagrams on this page specifically relating to the query.",
                            f"Query: {query}",
                            img_bytes, model, rpm
                        )
                        vision_info += f"\n[Vision Insight for Page {c['start_page']}]: {desc}\n"

        prompt = f"{m_ins}\nQuery: {query}\n" 
        if vision_info: prompt += f"\nADDITIONAL VISUAL CONTEXT:\n{vision_info}\n"
        prompt += "\n".join([f"ID:{c['id']} P:{c['start_page']} Text:{c['text']}" for c in batch])
        return await call_gemini_json_async(m_sys, prompt, model, rpm)

    results = await asyncio.gather(*[process_batch(b) for b in batches])
    out = []
    for r in results:
        if isinstance(r, list): out.extend(r)
        elif isinstance(r, dict): out.append(r)
    return [o for o in out if isinstance(o, dict) and "error" not in o]

async def recursive_reduce(items: List[Dict], query: str, model: str, rpm: int, mode: str):
    m_sys, m_ins, r_sys, r_ins = get_prompts(mode)
    
    if len(items) > 25:
        # Hierarchy: merge chunks into larger intermediate results
        sub_batches = [items[i:i+15] for i in range(0, len(items), 15)]
        async def summarize_sub(batch):
            prompt = f"Merge and deduplicate these partial {mode} entries for query '{query}'. Keep evidence/snippets intact. Output a JSON array in the format: {m_ins}\n\nData:\n{json.dumps(batch)}"
            # Use r_sys for intermediate reduction too, but specify format
            return await call_gemini_json_async(r_sys, prompt, model, rpm)
        
        intermediate = await asyncio.gather(*[summarize_sub(b) for b in sub_batches])
        items = []
        for it in intermediate:
            if isinstance(it, list): items.extend(it)
            elif isinstance(it, dict):
                if "error" not in it: items.append(it)
    
    user_prompt = f"{r_ins}\nQuery: {query}\n\nFindings Data:\n{json.dumps(items)}"
    return await call_gemini_json_async(r_sys, user_prompt, model, rpm)

# --------------------------
# UI: main interactions
# --------------------------
uploaded_files = st.file_uploader("Upload one or more tender PDFs", type=["pdf"], accept_multiple_files=True, key="file_uploader")
query_input = st.text_area("Enter one or more queries (newline separated)", height=120, key="query_text_area")

run_button = st.button("Run summarization", key="run_summarization_button")
bench_button = False
if "benchmark_df" in st.session_state:
    bench_button = st.button("🚀 Run Benchmark (from CSV)", key="run_benchmark_button")

if (uploaded_files and (query_input or "benchmark_df" in st.session_state) and (run_button or bench_button)):
    is_benchmark = bench_button
    queries_to_run = []
    ground_truths = {}
    
    if is_benchmark:
        df_bench = st.session_state["benchmark_df"]
        queries_to_run = df_bench["queries"].tolist()
        # Map query to ground truth answer
        for _, row in df_bench.iterrows():
            ground_truths[row["queries"]] = row["answer"]
    else:
        queries_to_run = [q.strip() for q in query_input.splitlines() if q.strip()]

    tmp_paths = []
    for uploaded_file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpf:
            tmpf.write(uploaded_file.getbuffer())
            tmp_paths.append(tmpf.name)

    st.info(f"Processing {len(tmp_paths)} uploaded PDF(s)...")

    all_pages = []
    all_image_flags = []
    total_extract_time = 0.0
    doc = None # Primary doc for rendering
    for p in tmp_paths:
        t0 = time.time()
        pages, image_flags = extract_pages(p)
        t1 = time.time()
        total_extract_time += (t1 - t0)
        basename = os.path.basename(p)
        pages = [f"[File {basename}] " + pg for pg in pages]
        all_pages.extend(pages)
        all_image_flags.extend(image_flags)
        if not doc: doc = fitz.open(p) # Use first file for citation rendering
    
    st.write(f"Extracted {len(all_pages)} pages from {len(tmp_paths)} file(s) (time: {total_extract_time:.2f}s).")

    t0 = time.time()
    chunks = semantic_chunk_pages(all_pages, all_image_flags, model, max_tokens=max_tokens_per_chunk, overlap_sentences=overlap_sentences)
    t1 = time.time()
    chunk_time = t1 - t0
    st.write(f"Created {len(chunks)} chunks (chunking time: {chunk_time:.2f}s).")
    
    if not chunks:
        st.error("No text could be extracted or chunked from the PDF. Please check if the PDF is readable/selectable.")
        st.stop()

    queries = queries_to_run
    results = []
    eval_results = []

    async def run_pipeline(doc: fitz.Document, enable_vision: bool):
        t_all_start = time.time()
        for q in queries:
            t0_bench_start = time.time()
            with st.status(f"Processing Query: {q}", expanded=True) as status:
                t0 = time.time()
                status.write("Running Map Phase (Scanning chunks)...")
                mapped = await map_phase_async(chunks, q, model, target_rpm, batch_size, analysis_mode, doc=doc, enable_vision=enable_vision)
                
                t1 = time.time()
                status.write(f"Map phase completed in {t1-t0:.2f}s (Found {len(mapped)} snippets)")
                
                if not mapped:
                    status.warning("Found 0 snippets. The model didn't find matched content.")
                    status.update(label=f"Query Finished (No hits): {q}", state="complete", expanded=False)
                    results.append({"query": q, "mode": analysis_mode, "result": {"summary": "No relevant info found."}})
                    continue
                
                t0 = time.time()
                status.write("Running Reduce Phase (Synthesizing results)...")
                reduced = await recursive_reduce(mapped, q, model, target_rpm, analysis_mode)
                t1 = time.time()
                status.write(f"Reduce phase completed in {t1-t0:.2f}s")
                status.update(label=f"Query Finished: {q}", state="complete", expanded=False)

            # Display results
            st.subheader(f"Results for: {q}")
            
            with st.expander("🔍 View Raw Snippets Found", expanded=False):
                if mapped: st.json(mapped)
                else: st.write("No snippets were extracted in the Map phase.")

            if "error" in reduced:
                st.error(f"Error in Reduce phase: {reduced['error']}")
            
            # Global citations for the query
            all_cites = []

            if analysis_mode == "Compliance Matrix":
                matrix_data = reduced.get("matrix", [])
                if matrix_data: 
                    st.table(matrix_data)
                    for m in matrix_data:
                        text = m.get("evidence", "")
                        c = m.get("citations") or m.get("page")
                        if isinstance(c, list): 
                            for pg in c: all_cites.append({"page": pg, "text": text})
                        elif isinstance(c, int): 
                            all_cites.append({"page": c, "text": text})
                elif "error" not in reduced: st.warning("No compliance items found.")
            
            elif analysis_mode == "Risk Assessment":
                risks = reduced.get("risks", [])
                if risks:
                    for risk in risks:
                        with st.expander(f"⚠️ {risk.get('clause')} ({risk.get('risk_level', 'High')})"):
                            st.write(f"**Reason**: {risk.get('reason')}")
                            text = risk.get("evidence", "")
                            if text: st.info(f"**Evidence**: *\"{text}\"*")
                            st.warning(f"**Mitigation**: {risk.get('mitigation', 'Check with legal team.')}")
                            c = risk.get("citations") or risk.get("page")
                            if isinstance(c, list):
                                for pg in c: all_cites.append({"page": pg, "text": text})
                            elif isinstance(c, int):
                                all_cites.append({"page": c, "text": text})
                elif "error" not in reduced: st.warning("No risks identified.")

            elif analysis_mode == "Entity Dashboard":
                dash = reduced.get("dashboard", {})
                if dash:
                    cols = st.columns(len(dash))
                    for i, (cat, items) in enumerate(dash.items()):
                        with cols[i]:
                            st.metric(label=cat, value=len(items))
                            for item in items:
                                st.markdown(f"**{item.get('entity', 'Item')}**: {item.get('value', 'N/A')}")
                                c = item.get("page")
                                if c: all_cites.append({"page": c, "text": item.get("value")})
                else: st.warning("Dashboard data could not be compiled.")

            elif analysis_mode == "Ambiguity Scrutiny":
                ambiguity_items = reduced.get("queries", [])
                if ambiguity_items:
                    for q_item in ambiguity_items:
                        with st.chat_message("assistant"):
                            st.markdown(f"**Issue**: {q_item.get('item')}")
                            st.info(f"**Proposed Query**: {q_item.get('query')}")
                            st.caption(f"Refers to: {q_item.get('conflict', 'Ambiguous clause')}")
                            c = q_item.get("citations") or q_item.get("page")
                            text = q_item.get("conflict", "")
                            if isinstance(c, list):
                                for pg in c: all_cites.append({"page": pg, "text": text})
                            elif isinstance(c, int):
                                all_cites.append({"page": c, "text": text})
                else: st.warning("No ambiguities detected.")

            else:
                summary_text = reduced.get("summary", "")
                if not summary_text and "_raw" in reduced:
                    summary_text = reduced["_raw"]
                
                st.markdown(summary_text or "No summary generated. Try adjusting the query or prompt instructions.")
                
                bullets = reduced.get("bullets", [])
                if bullets:
                    st.markdown("**Key Points:**")
                    for b in bullets: st.markdown(f"- {b}")
                
                if "citations" in reduced:
                    c = reduced["citations"]
                    if isinstance(c, list):
                        for pg in c: all_cites.append({"page": pg, "text": summary_text[:50]})
                    elif isinstance(c, int):
                        all_cites.append({"page": c, "text": summary_text[:50]})
            
            # Render context preview if citations exist
            if all_cites:
                render_citation_preview(doc, all_cites)

            results.append({"query": q, "mode": analysis_mode, "result": reduced})
            
            # --- EVALUATION BLOCK ---
            if is_benchmark and q in ground_truths:
                gold = str(ground_truths[q])
                # Extract plain text from reduced result
                pred = ""
                if isinstance(reduced, dict):
                    pred = reduced.get("summary", "") or reduced.get("_raw", "")
                    if not pred and "matrix" in reduced: pred = json.dumps(reduced["matrix"])
                
                p = token_precision(pred, gold)
                r = token_recall(pred, gold)
                f1 = 2*p*r/(p+r) if (p+r)>0 else 0.0
                rl = rouge_l(pred, gold)
                
                err_status = ""
                if "error" in mapped: err_status = "Map Err"
                elif "error" in reduced: err_status = reduced["error"][:20]

                eval_results.append({
                    "Query": q,
                    "Precision": round(p, 4),
                    "Recall": round(r, 4),
                    "F1": round(f1, 4),
                    "ROUGE-L": round(rl, 4),
                    "Time (s)": round(time.time() - t0_bench_start, 2),
                    "Status": err_status or "OK"
                })

        if is_benchmark and eval_results:
            st.markdown("---")
            st.header("📈 Benchmark Results")
            df_res = pd.DataFrame(eval_results)
            st.dataframe(df_res, use_container_width=True)
            
            # Aggregated Metrics
            cols = st.columns(4)
            cols[0].metric("Avg Precision", round(df_res["Precision"].mean(), 4))
            cols[1].metric("Avg Recall", round(df_res["Recall"].mean(), 4))
            cols[2].metric("Avg F1", round(df_res["F1"].mean(), 4))
            cols[3].metric("Avg ROUGE-L", round(df_res["ROUGE-L"].mean(), 4))

        st.success(f"Total processing completed in {time.time()-t_all_start:.2f}s")

    asyncio.run(run_pipeline(doc, enable_vision))

    st.markdown("---")
    st.write(f"Total extraction time: {total_extract_time:.2f}s")
    st.write(f"Chunking time: {chunk_time:.2f}s")
    
    total_processing_time = total_extract_time + chunk_time
    st.success(f"**Total processing time: {total_processing_time:.2f}s**")

    # Cleanup temp paths
    if doc: doc.close()
    for p in tmp_paths:
        try: os.unlink(p)
        except: pass

    out_txt = os.path.join(tempfile.gettempdir(), "tender_summary_merged.txt")
    out_json = os.path.join(tempfile.gettempdir(), "tender_summary_merged.json")
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write("TENDER ANALYSIS REPORT\n")
        f.write(f"Sources: {', '.join(os.path.basename(p) for p in tmp_paths)}\n\n")
        for r in results:
            mode = r["mode"]
            res = r["result"]
            f.write(f"### Query: {r['query']} ({mode})\n")
            if mode == "General Summary":
                f.write(f"\nSummary:\n{res.get('summary', 'N/A')}\n\nKey Points:\n")
                for b in res.get("bullets", []): f.write(f"- {b}\n")
            elif mode == "Compliance Matrix":
                f.write("\nCompliance Items:\n")
                for item in res.get("matrix", []):
                    f.write(f"- {item.get('item')}: {item.get('detail')}\n")
                    if item.get("evidence"): f.write(f"  Evidence: \"{item['evidence']}\"\n")
                    f.write(f"  (Citations: {item.get('citations') or item.get('page')})\n")
            elif mode == "Risk Assessment":
                f.write("\nIdentified Risks:\n")
                for risk in res.get("risks", []):
                    f.write(f"- {risk.get('clause')} ({risk.get('risk_level')}): {risk.get('reason')}\n")
                    if risk.get("evidence"): f.write(f"  Evidence: \"{risk['evidence']}\"\n")
                    f.write(f"  Mitigation: {risk.get('mitigation')}\n")
            elif mode == "Entity Dashboard":
                f.write("\nKey Entity Dashboard:\n")
                for cat, items in res.get("dashboard", {}).items():
                    f.write(f"\n[{cat}]\n")
                    for it in items: f.write(f"- {it.get('entity')}: {it.get('value')}\n")
            elif mode == "Ambiguity Scrutiny":
                f.write("\nAmbiguities & Clarifications:\n")
                for q_item in res.get("queries", []):
                    f.write(f"- Issue: {q_item.get('item')}\n")
                    f.write(f"  Suggested Query: {q_item.get('query')}\n")
            f.write("\n")
    with open(out_json, "w", encoding="utf-8") as jf:
        json.dump({"results": results}, jf, indent=2)

    with open(out_txt, "rb") as f_txt:
        st.download_button("Download TXT", data=f_txt.read(), file_name="tender_summary.txt")
    with open(out_json, "rb") as f_json:
        st.download_button("Download JSON", data=f_json.read(), file_name="tender_summary.json")

    # Workflow Acceleration: Clarification Letter
    doc_letter = generate_clarification_letter(results)
    if doc_letter:
        st.download_button(
            label="📝 Download Clarification Letter (.docx)",
            data=doc_letter,
            file_name="Pre_bid_Queries.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    # Structured Data Export: Excel
    xlsx_report = generate_excel_report(results)
    if xlsx_report:
        st.download_button(
            label="📊 Download Excel Report (.xlsx)",
            data=xlsx_report,
            file_name="Tender_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    st.markdown("---")
    st.write("Debug logs (last 10 items):")
    debug_files = sorted([os.path.join(DEBUG_DIR, f) for f in os.listdir(DEBUG_DIR)], key=os.path.getmtime, reverse=True)[:10]
    for p in debug_files:
        with st.expander(os.path.basename(p)):
            try:
                with open(p, "r", encoding="utf-8") as fh:
                    st.code(fh.read()[:5000])
            except Exception:
                st.write("Error reading file.")

elif not uploaded_files:
    st.info("Upload tender PDF(s) to begin.")
    st.stop()
elif not query_input:
    st.info("Enter one or more queries.")
    st.stop()
